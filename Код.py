# Імпорт необхідних бібліотек
import tkinter as tk
from tkinter import messagebox
import docx
import os

# Функція для генерації документа Word
def generate_document(data):
    doc = docx.Document()
    doc.add_heading('Документ', level=1)

    if data["option"] == "Скасовуємо":
        doc.add_paragraph(f"Цей документ скасовує рішення. Заявник: {data['applicant']}.")
    else:
        doc.add_paragraph(f"Цей документ залишає рішення в силі. Заявник: {data['applicant']}.")

    doc.add_paragraph(f"Рік народження: {data['birth_year']}.")
    doc.add_paragraph(f"Номер постанови: {data['decision_number']}.")

    if data["email"]:  # Якщо поле електронної пошти не порожнє
        doc.add_paragraph(f"Електронна пошта: {data['email']}.")

    # Визначення шляху для збереження файлу (сумісно для MacOS та Windows)
    home_path = os.path.expanduser('~')  # Отримання домашньої директорії
    desktop = os.path.join(home_path, 'Desktop')
    filepath = os.path.join(desktop, 'output.docx')

    try:
        doc.save(filepath)
        messagebox.showinfo("Інформація", f"Документ сформовано!\nЗбережено в: {filepath}")
    except PermissionError:
        messagebox.showerror("Помилка", "Немає прав на запис у вказану директорію.")

# Функція, що викликається при натисканні на кнопку в GUI
def submit():
    # Збирання даних із полів вводу
    data = {
        "option": option_var.get(),
        "applicant": entry_applicant.get(),
        "birth_year": entry_birth_year.get(),
        "decision_number": entry_decision_number.get(),
        "email": entry_email.get()  # Збір даних з нового поля
        # Додайте тут інші поля
    }
    generate_document(data)  # Виклик функції генерації документа з зібраними даними

# Створення головного вікна програми
window = tk.Tk()
window.title("Генератор Документів")

# Створення та розміщення радіокнопок для вибору опцій
option_var = tk.StringVar()
option_var.set("Залишаємо в силі")  # Встановлення вихідного значення

tk.Radiobutton(window, text="Залишаємо в силі", variable=option_var, value="Залишаємо в силі").pack()
tk.Radiobutton(window, text="Скасовуємо", variable=option_var, value="Скасовуємо").pack()

# Створення та розміщення полів для введення тексту
tk.Label(window, text="Від кого (заявник):").pack()
entry_applicant = tk.Entry(window)
entry_applicant.pack()

tk.Label(window, text="Рік народження:").pack()
entry_birth_year = tk.Entry(window)
entry_birth_year.pack()

tk.Label(window, text="Номер постанови:").pack()
entry_decision_number = tk.Entry(window)
entry_decision_number.pack()

tk.Label(window, text="Електронна пошта:").pack()
entry_email = tk.Entry(window)
entry_email.pack()

# Додайте тут інші поля для введення даних

# Створення та розміщення кнопки для сформування документа
submit_button = tk.Button(window, text="Сформувати Документ", command=submit)
submit_button.pack()

# Запуск головного циклу обробки подій GUI
window.mainloop()
